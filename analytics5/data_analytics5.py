import pandas as pd
import sqlite3
import plotly.express as px
from scipy.stats import chi2_contingency
from jinja2 import Environment
from docx import Document
from weasyprint import HTML
import base64
import os

# === CONFIG ===
EXCEL_FILE = "data.xlsx"
DB_FILE = "harassment_survey.db"
CHART_DIR = "charts"
REPORT_HTML = "harassment_policy_report.html"
REPORT_PDF = "harassment_policy_report.pdf"
REPORT_DOCX = "executive_summary.docx"
EMBED_XLSX = "survey_data.xlsx"

# === LOAD DATA ===
df = pd.read_excel(EXCEL_FILE)
df.columns = [c.strip().replace(" ", "_").replace("?", "").replace(",", "").replace("/", "_") for c in df.columns]
df.to_excel(EMBED_XLSX, index=False)

# Base64-encoded original data for embedding
with open(EMBED_XLSX, "rb") as f:
    encoded_excel = base64.b64encode(f.read()).decode()

os.makedirs(CHART_DIR, exist_ok=True)

# === CONVERT TO SQLITE ===
conn = sqlite3.connect(DB_FILE)
df.to_sql("survey", conn, if_exists="replace", index=False)

# === ANALYSIS & CHARTS ===

def plot_and_save(df, title, filename, x, y='Count', color=None):
    fig = px.bar(df, x=x, y=y, color=color, barmode="group", text_auto=True, title=title)
    chart_path = os.path.join(CHART_DIR, filename)
    fig.write_html(chart_path)
    return chart_path

def chi_square_test(df, col1, col2):
    contingency = pd.crosstab(df[col1], df[col2])
    
    # Check if contingency table is empty
    if contingency.empty:
        return f"Chi-square test between '{col1}' and '{col2}': No data available for chi-square test.", None
    
    chi2, p, dof, expected = chi2_contingency(contingency)
    return f"Chi-square test between '{col1}' and '{col2}': χ² = {chi2:.2f}, p = {p:.4f} (dof={dof})", p

# 1. Awareness Levels
awareness_df = df['Are_you_aware_of_what_constitutes_workplace_harassment'].value_counts().reset_index()
awareness_df.columns = ['Awareness', 'Count']
chart_awareness = plot_and_save(awareness_df, "Awareness of Workplace Harassment", "awareness.html", x='Awareness')

# 2. Awareness vs Reporting
awareness_reporting_df = df.groupby(['Are_you_aware_of_what_constitutes_workplace_harassment',
                                     'Do_you_know_whom_to_report_workplace_harassment_incidents_to']).size().reset_index(name='Count')
awareness_reporting_df.columns = ['Awareness', 'Know_Whom', 'Count']
chart_awareness_reporting = plot_and_save(awareness_reporting_df, "Awareness vs Reporting Knowledge", "awareness_reporting.html", x='Awareness', color='Know_Whom')
chi_awareness_reporting, p1 = chi_square_test(df, 'Are_you_aware_of_what_constitutes_workplace_harassment',
                                               'Do_you_know_whom_to_report_workplace_harassment_incidents_to')

# 3. Training vs Reporting
training_reporting_df = df.groupby(['Have_you_ever_received_any_formal_training_on_workplace_harassment_policies',
                                    'Do_you_know_whom_to_report_workplace_harassment_incidents_to']).size().reset_index(name='Count')
training_reporting_df.columns = ['Training', 'Know_Whom', 'Count']

# Remove rows with missing data in columns for chi-square test
training_reporting_df = training_reporting_df.dropna(subset=['Training', 'Know_Whom'])

chart_training_reporting = plot_and_save(training_reporting_df, "Training vs Reporting Knowledge", "training_reporting.html", x='Training', color='Know_Whom')
chi_training_reporting, p2 = chi_square_test(training_reporting_df, 
                                              'Training',
                                              'Know_Whom')

# 4. Gender-based Awareness
gender_awareness_df = df.groupby(['What_is_your_Gender',
                                  'Are_you_aware_of_what_constitutes_workplace_harassment']).size().reset_index(name='Count')
gender_awareness_df.columns = ['Gender', 'Awareness', 'Count']
chart_gender_awareness = plot_and_save(gender_awareness_df, "Gender-based Awareness", "gender_awareness.html", x='Gender', color='Awareness')
chi_gender_awareness, p3 = chi_square_test(df, 'What_is_your_Gender',
                                           'Are_you_aware_of_what_constitutes_workplace_harassment')

# === INTERPRETATION ===
summary = f"""
<strong>Executive Summary</strong><br>
Of the {len(df)} respondents, {awareness_df[awareness_df['Awareness'] == 'Yes']['Count'].values[0]} ({awareness_df[awareness_df['Awareness'] == 'Yes']['Count'].values[0]/len(df)*100:.1f}%) are aware of what constitutes workplace harassment.<br>
{chi_awareness_reporting}<br>
{chi_training_reporting}<br>
{chi_gender_awareness}
"""

recommendations = """
<ul>
  <li><strong>Mandatory Training:</strong> Formal training should be instituted across departments to improve awareness and reporting clarity.</li>
  <li><strong>Clarify Reporting Channels:</strong> Low awareness of whom to report to indicates a need for clearer communication protocols.</li>
  <li><strong>Target Gender Disparities:</strong> Gender-based differences in awareness require tailored outreach.</li>
  <li><strong>Continuous Learning:</strong> Refresher programs should be scheduled periodically.</li>
  <li><strong>Accessible Policy Documents:</strong> Ensure policies are visible and easy to locate on internal platforms.</li>
</ul>
"""

# === GENERATE HTML REPORT ===
html_template = f"""
<!DOCTYPE html>
<html>
<head>
  <title>Workplace Harassment Policy Report</title>
  <style>
    body {{ font-family: Arial; margin: 40px; background: #f4f4f4; }}
    .section {{ background: #fff; padding: 20px; border-radius: 10px; margin-bottom: 30px; }}
    h1, h2 {{ color: #2c3e50; }}
    iframe {{ width: 100%; height: 500px; border: none; }}
  </style>
</head>
<body>
  <h1>Comprehensive Report on Workplace Harassment Policy Awareness</h1>
  <div class="section"><h2>1. Executive Summary</h2><p>{summary}</p></div>

  <div class="section"><h2>2. Awareness Levels</h2><iframe src="{chart_awareness}"></iframe></div>
  <div class="section"><h2>3. Awareness vs Reporting</h2><iframe src="{chart_awareness_reporting}"></iframe></div>
  <div class="section"><h2>4. Training vs Reporting</h2><iframe src="{chart_training_reporting}"></iframe></div>
  <div class="section"><h2>5. Gender-based Awareness</h2><iframe src="{chart_gender_awareness}"></iframe></div>

  <div class="section"><h2>6. Download Dataset</h2>
    <a download="SurveyData.xlsx" href="data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{encoded_excel}">Download Excel File</a>
  </div>

  <div class="section"><h2>7. Recommendations</h2>{recommendations}</div>
</body>
</html>
"""

with open(REPORT_HTML, "w", encoding="utf-8") as f:
    f.write(html_template)

# === EXPORT TO PDF ===
HTML(REPORT_HTML).write_pdf(REPORT_PDF)

# === EXPORT TO DOCX ===
doc = Document()
doc.add_heading("Executive Summary – Workplace Harassment Survey", level=1)
doc.add_paragraph(summary.replace("<strong>", "").replace("</strong>", "").replace("<br>", "\n"))
doc.add_heading("Recommendations", level=2)
doc.add_paragraph("""
1. Mandatory Training: Formal training on harassment policies should be made mandatory for all employees.
2. Clear Reporting Channels: Ensure all employees know whom to contact in case of harassment.
3. Gender-Sensitive Communication: Address awareness gaps by tailoring outreach strategies.
4. Continuous Learning: Reinforce knowledge with scheduled refresher courses.
5. Visible Policy Access: Make policy documents easy to access for all staff.
""")
doc.save(REPORT_DOCX)

print(f"✅ HTML report generated: {REPORT_HTML}")
print(f"✅ PDF exported: {REPORT_PDF}")
print(f"✅ Word summary saved: {REPORT_DOCX}")

